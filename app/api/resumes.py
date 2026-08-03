"""
app/api/resumes.py
===================
简历相关 API。

POST /resumes/upload         → 上传简历文本，创建候选人记录
POST /resumes/upload-file    → 上传 PDF/DOCX/TXT 文件，自动提取文本
POST /resumes/{id}/parse     → 调用 Resume Agent 解析
GET  /resumes/{id}           → 获取候选人详情
GET  /resumes/               → 获取所有候选人列表
"""

import uuid
import tempfile
import os
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from pydantic import BaseModel, Field

from app.database.session import get_db
from app.database import crud
from app.agents.resume_agent import parse_resume

router = APIRouter(prefix="/resumes", tags=["resumes"])


# ---- 请求/响应模型 ----

class ResumeUploadRequest(BaseModel):
    """上传简历文本的请求体。"""
    resume_text: str
    name: str | None = None
    filename: str | None = None


class CandidateNameUpdateRequest(BaseModel):
    """人工修改候选人姓名的请求体。"""

    # 限制长度可以拦住空字符串和误粘贴的整段简历。
    name: str = Field(..., min_length=1, max_length=80)


# ---- 辅助: 自动生成申请人名称 ----
# 当简历中没有姓名时 (如匿名化简历)，自动分配 "申请人A", "申请人B" ...
# 用 A-Z 循环: 申请人A...申请人Z, 申请人AA, 申请人AB...
def _generate_applicant_name(db: Session) -> str:
    """
    根据数据库中已有候选人数量生成名称。

    规则: 申请人A, 申请人B, ..., 申请人Z, 申请人AA, 申请人AB, ...
    确保即使有大量匿名简历也不会重名。

    参数:
        db: 数据库会话
    返回:
        str: 如 "申请人A", "申请人B"
    """
    # 查当前候选人总数
    count = db.query(crud.models.Candidate).count()
    # 转字母: 0→A, 1→B, ..., 25→Z, 26→AA, 27→AB...
    def num_to_letters(n: int) -> str:
        result = ""
        while n >= 0:
            result = chr(ord('A') + n % 26) + result
            n = n // 26 - 1
        return result
    suffix = num_to_letters(count)
    return f"申请人{suffix}"


def _extract_file_text(filename: str, content: bytes) -> str:
    """
    从上传的文件中提取纯文本。

    根据扩展名自动选择解析器:
    - .pdf  → PyMuPDF (fitz)
    - .docx → python-docx
    - .txt  → 直接解码

    参数:
        filename: 原始文件名 (用于判断格式)
        content: 文件二进制内容
    返回:
        str: 提取的纯文本
    """
    import io

    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        # PDF: 通过 PyMuPDF 读取
        import fitz  # PyMuPDF
        text_parts = []
        # fitz.open 可以接受 bytes 流 (通过 stream 参数)
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc:
            # sort=True 会按页面坐标从上到下、从左到右整理文本。
            # PDF 内部对象顺序经常和视觉顺序不同，不排序时姓名可能跑到“个人概述”之后。
            text_parts.append(page.get_text("text", sort=True))
        doc.close()
        return "\n".join(text_parts)

    elif ext == "docx":
        # DOCX: python-docx 读取
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)

    elif ext == "txt" or ext == "md":
        # TXT/MD: 直接解码
        return content.decode("utf-8", errors="replace")

    else:
        raise ValueError(f"不支持的文件格式: .{ext}，支持的格式: PDF, DOCX, TXT, MD")


# ---- API 端点 ----

@router.post("/upload")
async def upload_resume(
    request: ResumeUploadRequest,
    db: Session = Depends(get_db),
):
    """
    上传简历文本，创建候选人记录。

    请求体:
      {"resume_text": "简历全文...", "name": "可选", "filename": "简历.pdf"}
    """
    # 文本粘贴和文件上传使用同一套自动命名，后续解析到可信姓名时才允许覆盖。
    display_name = request.name.strip() if request.name and request.name.strip() else _generate_applicant_name(db)
    candidate = crud.create_candidate(
        db=db,
        resume_text=request.resume_text,
        name=display_name,
        filename=request.filename,
    )
    return {
        "candidate_id": candidate.candidate_id,
        "name": candidate.name,
        "message": "简历上传成功，请调用 /resumes/{id}/parse 进行解析",
    }


@router.post("/upload-file")
async def upload_resume_file(
    file: UploadFile = File(...),
    name: str | None = Form(None),
    db: Session = Depends(get_db),
):
    """
    上传 PDF/DOCX/TXT 简历文件，自动提取文本并创建候选人。

    支持格式: PDF, DOCX, TXT, MD
    如未提供姓名，自动分配"申请人A/B/C..."。

    请求: multipart/form-data
      - file: 简历文件 (必填)
      - name: 候选人姓名 (可选，匿名简历自动命名)

    返回: candidate_id + 提取的文本
    """
    # Step 1: 校验文件格式
    filename = file.filename or "unknown"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    allowed = {"pdf", "docx", "txt", "md"}
    if ext not in allowed:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式 .{ext}，支持的格式: {', '.join(allowed)}",
        )

    # Step 2: 读取文件内容
    rag_indexed = False
    rag_index_warning = None
    try:
        content = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件读取失败: {str(e)}")

    if not content:
        raise HTTPException(status_code=400, detail="文件为空")

    # Step 3: 提取文本
    try:
        resume_text = _extract_file_text(filename, content)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"文件解析失败: {str(e)}。请确认文件未损坏且格式正确。",
        )

    if not resume_text.strip():
        raise HTTPException(
            status_code=400,
            detail="文件中未提取到文本内容。请确认PDF不是扫描件/图片。",
        )

    # Step 4: 自动命名 (匿名简历)
    display_name = name if name and name.strip() else _generate_applicant_name(db)

    # Step 5: 存入数据库
    candidate = crud.create_candidate(
        db=db,
        resume_text=resume_text,
        name=display_name,
        filename=filename,
    )

    return {
        "candidate_id": candidate.candidate_id,
        "name": candidate.name,
        "filename": filename,
        "text_length": len(resume_text),
        "text_preview": resume_text[:200] + "..." if len(resume_text) > 200 else resume_text,
        "message": f"文件上传成功 ({len(resume_text)} 字符)。请调用 /resumes/{candidate.candidate_id}/parse 进行解析",
    }


@router.post("/{candidate_id}/parse")
async def parse_resume_endpoint(
    candidate_id: str,
    db: Session = Depends(get_db),
):
    """
    调用 Resume Agent 解析简历。
    """
    candidate = crud.get_candidate(db, candidate_id)
    if not candidate:
        raise HTTPException(status_code=404, detail="候选人不存在")

    try:
        profile = await parse_resume(
            resume_text=candidate.resume_text,
            candidate_id=candidate_id,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"简历解析失败: {str(e)}")

    crud.update_candidate_profile(db, candidate_id, profile)

    # ---- RAG 索引: 切分简历 → Embedding → Qdrant ----
    from app.services.rag_service import index_resume_text
    try:
        point_ids = index_resume_text(
            resume_text=candidate.resume_text,
            candidate_id=candidate_id,
            source="resume_parse",
        )
        # 保存 chunk → Qdrant point 映射到数据库
        from app.services.document_loader import chunk_documents
        from langchain_core.documents import Document
        doc = Document(page_content=candidate.resume_text, metadata={"source": "resume_parse"})
        chunks = chunk_documents([doc])
        crud.save_resume_chunks(
            db,
            candidate_id,
            chunks,
            point_ids,
            replace_existing=True,
        )
        rag_indexed = bool(point_ids)
    except Exception as error:
        # 画像解析仍然保留，但必须把 RAG 失败明确返回，不能再静默伪装成解析全成功。
        rag_index_warning = (
            "候选人画像已解析，但简历证据索引失败；请确认 Qdrant 已启动且 "
            f"Embedding 模型可用。原因：{str(error)[:200]}"
        )

    return {
        "candidate_id": candidate_id,
        "profile": profile,
        "rag_indexed": rag_indexed,
        "rag_index_warning": rag_index_warning,
    }


@router.patch("/{candidate_id}")
async def update_candidate_name(
    candidate_id: str,
    request: CandidateNameUpdateRequest,
    db: Session = Depends(get_db),
):
    """人工修改候选人姓名，并同步更新结构化画像中的姓名。"""
    # Pydantic 会检查长度；这里再去掉用户无意输入的首尾空格。
    clean_name = request.name.strip()
    if not clean_name:
        raise HTTPException(status_code=400, detail="候选人姓名不能为空")

    candidate = crud.update_candidate_name(db, candidate_id, clean_name)
    if not candidate:
        raise HTTPException(status_code=404, detail="候选人不存在")

    return {
        "candidate_id": candidate.candidate_id,
        "name": candidate.name,
        "profile": candidate.profile_json,
        "message": "候选人姓名已更新",
    }


@router.get("/{candidate_id}")
async def get_candidate(
    candidate_id: str,
    db: Session = Depends(get_db),
):
    """获取候选人详情。"""
    candidate = crud.get_candidate(db, candidate_id)
    if not candidate:
        raise HTTPException(status_code=404, detail="候选人不存在")

    return {
        "candidate_id": candidate.candidate_id,
        "name": candidate.name,
        "email": candidate.email,
        "resume_filename": candidate.resume_filename,
        "resume_text": candidate.resume_text,
        "profile": candidate.profile_json,
    }


@router.get("/")
async def list_candidates(db: Session = Depends(get_db)):
    """获取所有候选人列表。"""
    candidates = crud.get_all_candidates(db)
    return [
        {
            "candidate_id": c.candidate_id,
            "name": c.name,
            "email": c.email,
            "filename": c.resume_filename,
            "has_profile": c.profile_json is not None,
        }
        for c in candidates
    ]


@router.delete("/{candidate_id}")
async def delete_candidate_endpoint(candidate_id: str, db: Session = Depends(get_db)):
    """删除候选人及其关联的简历chunk和匹配结果。"""
    if crud.delete_candidate(db, candidate_id):
        return {"message": "候选人已删除", "candidate_id": candidate_id}
    raise HTTPException(status_code=404, detail="候选人不存在")
