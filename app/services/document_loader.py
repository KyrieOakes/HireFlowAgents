"""
app/services/document_loader.py
=================================
文档加载服务。

使用 LangChain Document Loaders 加载各种格式的文件:
- PDF: PyMuPDFLoader (后端用 PyMuPDF/fitz)
- DOCX: python-docx (手动处理, LangChain 的 DocxLoader 有限)
- TXT: 直接读取

加载后用 RecursiveCharacterTextSplitter 切分为 chunks:
- 按 \\n\\n → \\n → 空格 → 字符的顺序递归尝试切分
- 尽量在段落/句子边界处切分，保持语义完整性
- chunk_size=500, chunk_overlap=50 (可从配置调整)

依赖:
- PyMuPDF (fitz): pip install PyMuPDF
- python-docx: pip install python-docx
- langchain-text-splitters: pip install langchain-text-splitters
"""

from typing import List
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from app.utils.config import settings


# ============================================================
# 文档加载器
# ============================================================

def load_pdf(file_path: str) -> List[Document]:
    """
    加载 PDF 文件，返回 LangChain Document 对象列表。

    每一页 PDF 会被转换为一个 Document。
    Document 包含:
    - page_content: 该页的纯文本内容
    - metadata: {"source": 文件路径, "page": 页码}

    参数:
        file_path: PDF 文件路径
    返回:
        List[Document]: 每页一个 Document 对象
    """
    # 使用 PyMuPDFLoader，它内部调用 fitz (PyMuPDF) 提取文本
    # PyMuPDF 是 PDF 解析领域最快、最准确的库之一
    from langchain_community.document_loaders import PyMuPDFLoader

    loader = PyMuPDFLoader(file_path)
    # load() 返回 List[Document]，每个 Document 是 PDF 的一页
    documents = loader.load()
    return documents


def load_docx(file_path: str) -> List[Document]:
    """
    加载 DOCX 文件，返回 LangChain Document 对象列表。

    python-docx 提取纯文本后封装为单个 Document。
    因为 DOCX 没有"页码"概念，整个文档作为一个 Document。

    参数:
        file_path: DOCX 文件路径
    返回:
        List[Document]: 包含整个文档文本的单个 Document 列表
    """
    # python-docx 可以读取 Word 文档的段落和表格
    import docx

    doc = docx.Document(file_path)

    # 遍历所有段落，提取文本
    # 每个 paragraph 对象代表文档中的一个段落
    full_text = []
    for paragraph in doc.paragraphs:
        # paragraph.text 是段落的纯文本 (去除了格式信息)
        full_text.append(paragraph.text)

    # 用换行符连接所有段落
    text = "\n".join(full_text)

    # 封装为 LangChain Document 对象
    # metadata 中记录文件来源
    document = Document(
        page_content=text,
        metadata={"source": file_path},
    )

    return [document]


def load_txt(file_path: str) -> List[Document]:
    """
    加载 TXT 文件，返回 LangChain Document 对象列表。

    参数:
        file_path: TXT 文件路径
    返回:
        List[Document]: 包含整个文件文本的单个 Document 列表
    """
    # UTF-8 是最常见的文本编码
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    document = Document(
        page_content=text,
        metadata={"source": file_path},
    )

    return [document]


# ============================================================
# 文档加载入口 (自动识别格式)
# ============================================================

def load_document(file_path: str) -> List[Document]:
    """
    自动识别文件格式并加载文档。

    根据文件扩展名分派到对应的加载器。
    支持的格式: .pdf, .docx, .txt

    参数:
        file_path: 文件路径
    返回:
        List[Document]: 加载后的 Document 对象列表
    抛出:
        ValueError: 不支持的文件格式
    """
    # str.lower() 转小写，方便比较
    # str.endswith() 检查字符串是否以指定后缀结尾
    file_path_lower = file_path.lower()

    if file_path_lower.endswith(".pdf"):
        return load_pdf(file_path)
    elif file_path_lower.endswith(".docx"):
        return load_docx(file_path)
    elif file_path_lower.endswith(".txt"):
        return load_txt(file_path)
    else:
        # 不支持的文件格式，抛出异常
        raise ValueError(f"不支持的文件格式: {file_path}，支持的格式: PDF, DOCX, TXT")


# ============================================================
# 文本切分器
# ============================================================

def _get_text_splitter() -> RecursiveCharacterTextSplitter:
    """
    创建 RecursiveCharacterTextSplitter 实例。

    这个切分器的核心逻辑:
    1. 先尝试用 "\n\n" (段落分隔符) 切分
    2. 如果 chunk 还是太大，用 "\n" (换行符) 切分
    3. 如果还是太大，用 " " (空格/词边界) 切分
    4. 最后才用空字符串 "" (逐字符) 切分

    这种递归策略尽量让每个 chunk 在语义上完整，
    而不是在句子中间生硬切断。

    返回:
        RecursiveCharacterTextSplitter: 配置好的切分器
    """
    doc_config = settings.document

    return RecursiveCharacterTextSplitter(
        # chunk_size: 每个文本块的目标大小 (字符数)
        chunk_size=doc_config.chunk_size,
        # chunk_overlap: 相邻块之间的重叠字符数
        # 重叠可以防止关键信息正好落在两个 chunk 的边界处被"撕裂"
        chunk_overlap=doc_config.chunk_overlap,
        # 分隔符列表: 按优先级从高到低排列
        # 切分器会先用第一个分隔符尝试，切不动再用下一个
        separators=["\n\n", "\n", "。", "，", " ", ""],
        # length_function: 用字符数 (len) 来衡量 chunk 大小
        length_function=len,
    )


def chunk_documents(documents: List[Document]) -> List[Document]:
    """
    将 Document 列表切分为更小的 chunks。

    切分的原因:
    1. LLM 的上下文窗口有长度限制，不能一次塞入整个文档
    2. 向量检索时，更小的 chunk 意味着更精确的匹配
       (一个 chunk 只包含一个主题，embedding 更能代表它的含义)

    参数:
        documents: 要切分的 Document 列表
    返回:
        List[Document]: 切分后的 Document 列表
    """
    splitter = _get_text_splitter()

    # split_documents() 递归切分每个 Document
    # 并保留原始 metadata (如 source, page 等)
    chunks = splitter.split_documents(documents)

    return chunks
